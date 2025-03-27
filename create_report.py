from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import RGBColor

def create_title(doc, text, size=16, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    paragraph.alignment = alignment
    return paragraph

def add_content(doc, text, size=12, bold=False, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    paragraph.alignment = alignment
    return paragraph

def create_report():
    doc = Document()
    
    # Title Page
    doc.add_paragraph()  # Space for college logo
    create_title(doc, "DEPARTMENT OF COMPUTER SCIENCE AND ENGINEERING")
    create_title(doc, "[College Name]")
    create_title(doc, "[College Address]")
    doc.add_paragraph()
    create_title(doc, "A PROJECT REPORT")
    create_title(doc, "ON")
    create_title(doc, '"COLLEGE ADMISSION PORTAL"')
    create_title(doc, "A Web-Based Application for College Admissions")
    doc.add_paragraph()
    create_title(doc, "Submitted in partial fulfillment of the requirements for the award of the degree of")
    create_title(doc, "BACHELOR OF TECHNOLOGY")
    create_title(doc, "IN")
    create_title(doc, "COMPUTER SCIENCE AND ENGINEERING")
    doc.add_paragraph()
    create_title(doc, "Submitted by:")
    create_title(doc, "[Your Name]")
    create_title(doc, "[Your Roll Number]")
    doc.add_paragraph()
    create_title(doc, "Under the guidance of")
    create_title(doc, "[Professor Name]")
    create_title(doc, "[Department Name]")
    doc.add_page_break()

    # Certificate
    create_title(doc, "CERTIFICATE")
    doc.add_paragraph()
    add_content(doc, """This is to certify that the project work entitled "COLLEGE ADMISSION PORTAL" is a bonafide work carried out by [Your Name] (Roll No: [Your Roll Number]) under my guidance and supervision for the partial fulfillment of the requirements for the award of the degree of Bachelor of Technology in Computer Science and Engineering.""")
    doc.add_paragraph()
    doc.add_paragraph()
    add_content(doc, "[Professor Name]")
    add_content(doc, "[Department Name]")
    add_content(doc, "[College Name]")
    doc.add_paragraph()
    add_content(doc, "[Date]")
    doc.add_page_break()

    # Problem Statement
    create_title(doc, "1. Problem Statement", size=14)
    doc.add_paragraph()
    create_title(doc, "1.1. Problem Statement Sr. No: 7", size=12)
    create_title(doc, "1.2. Project Title: College Admission Portal", size=12)
    create_title(doc, "1.3. Problem Statement:", size=12)
    add_content(doc, """Educational institutions require a web-based system to manage student admissions efficiently. The system needs to handle online application submissions, document uploads, and application status tracking. It should provide role-based access for students and administrators, with features for application processing, status updates, and secure data management.""")
    doc.add_page_break()

    # Introduction
    create_title(doc, "2. Introduction", size=14)
    add_content(doc, """The College Admission Portal is a web-based Java application designed to modernize and streamline the college admission process. This application provides a user-friendly interface for students to submit applications online and for administrators to manage the admission process. The system is built using Java Servlets and JSP for the web interface and implements MVC architecture for efficient data management and processing.""")
    doc.add_paragraph()

    # Methodology
    create_title(doc, "3. Methodology", size=14)
    add_content(doc, "The project follows a structured development approach:")
    add_content(doc, "1. Analysis Phase:")
    add_content(doc, "   • Requirements gathering")
    add_content(doc, "   • System architecture design")
    add_content(doc, "   • Database design planning")
    add_content(doc, "2. Design Phase:")
    add_content(doc, "   • Class design (User, Application, Document)")
    add_content(doc, "   • Web interface design")
    add_content(doc, "   • Database schema design")
    add_content(doc, "3. Implementation Phase:")
    add_content(doc, "   • Frontend development using JSP and Bootstrap")
    add_content(doc, "   • Backend development using Java Servlets")
    add_content(doc, "   • Database implementation using MySQL")
    add_content(doc, "   • Testing and deployment")
    doc.add_page_break()

    # Implementation and Results
    create_title(doc, "4. Implementation and Results", size=14)
    add_content(doc, "The implementation includes:")
    add_content(doc, "1. Core Classes:")
    
    # Add code snippets
    add_content(doc, "a) User.java:")
    doc.add_paragraph("```java\npublic class User {\n    private int id;\n    private String email;\n    // ... rest of the code\n}```")
    
    # Add more sections...
    
    # Save the document
    doc.save('College_Admission_Portal_Report.docx')

if __name__ == "__main__":
    create_report()
    print("Report has been generated successfully!") 