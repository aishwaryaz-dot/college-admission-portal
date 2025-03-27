from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import RGBColor

def create_title(doc, text, size=16, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    paragraph.alignment = alignment
    return paragraph

def add_content(doc, text, size=12, bold=False, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    paragraph.alignment = alignment
    return paragraph

def add_code(doc, code, language="java"):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(f"```{language}\n{code}\n```")
    run.font.size = Pt(10)
    run.font.name = 'Courier New'
    return paragraph

def create_report():
    doc = Document()
    
    # Title Page
    doc.add_paragraph()  # Space for college logo
    create_title(doc, "DEPARTMENT OF COMPUTER SCIENCE AND ENGINEERING")
    create_title(doc, "[College Name]")
    create_title(doc, "[College Address]")
    doc.add_paragraph()
    create_title(doc, "A PROJECT REPORT")
    create_title(doc, "ON")
    create_title(doc, '"COLLEGE ADMISSION PORTAL"')
    create_title(doc, "A Web-Based Application for College Admissions")
    doc.add_paragraph()
    create_title(doc, "Submitted in partial fulfillment of the requirements for the award of the degree of")
    create_title(doc, "BACHELOR OF TECHNOLOGY")
    create_title(doc, "IN")
    create_title(doc, "COMPUTER SCIENCE AND ENGINEERING")
    doc.add_paragraph()
    create_title(doc, "Submitted by:")
    create_title(doc, "[Your Name]")
    create_title(doc, "[Your Roll Number]")
    doc.add_paragraph()
    create_title(doc, "Under the guidance of")
    create_title(doc, "[Professor Name]")
    create_title(doc, "[Department Name]")
    doc.add_page_break()

    # Certificate
    create_title(doc, "CERTIFICATE")
    doc.add_paragraph()
    add_content(doc, """This is to certify that the project work entitled "COLLEGE ADMISSION PORTAL" is a bonafide work carried out by [Your Name] (Roll No: [Your Roll Number]) under my guidance and supervision for the partial fulfillment of the requirements for the award of the degree of Bachelor of Technology in Computer Science and Engineering.""")
    doc.add_paragraph()
    doc.add_paragraph()
    add_content(doc, "[Professor Name]")
    add_content(doc, "[Department Name]")
    add_content(doc, "[College Name]")
    doc.add_paragraph()
    add_content(doc, "[Date]")
    doc.add_page_break()

    # Acknowledgment
    create_title(doc, "ACKNOWLEDGMENT")
    doc.add_paragraph()
    add_content(doc, """I would like to express my sincere gratitude to all those who have helped me in completing this project successfully. First and foremost, I would like to thank [Professor Name] for their valuable guidance and support throughout the project. Their expertise and constructive feedback have been instrumental in shaping this project.

I would also like to thank the Department of Computer Science and Engineering for providing the necessary resources and infrastructure. Special thanks to all the faculty members who have contributed to my learning and development.

Last but not least, I would like to thank my family and friends for their constant encouragement and support throughout this journey.""")
    doc.add_paragraph()
    add_content(doc, "[Your Name]")
    add_content(doc, "[Your Roll Number]")
    doc.add_page_break()

    # Declaration
    create_title(doc, "DECLARATION")
    doc.add_paragraph()
    add_content(doc, """I hereby declare that the project work entitled "COLLEGE ADMISSION PORTAL" submitted for the partial fulfillment of the requirements for the award of the degree of Bachelor of Technology in Computer Science and Engineering is my original work and has not been submitted for the award of any other degree, diploma, fellowship, or any other similar title or prize.""")
    doc.add_paragraph()
    add_content(doc, "[Your Name]")
    add_content(doc, "[Your Roll Number]")
    doc.add_paragraph()
    add_content(doc, "[Date]")
    doc.add_page_break()

    # Problem Statement
    create_title(doc, "1. Problem Statement", size=14)
    doc.add_paragraph()
    create_title(doc, "1.1. Problem Statement Sr. No: 7", size=12)
    create_title(doc, "1.2. Project Title: College Admission Portal", size=12)
    create_title(doc, "1.3. Problem Statement:", size=12)
    add_content(doc, """Educational institutions require a web-based system to manage student admissions efficiently. The system needs to handle online application submissions, document uploads, and application status tracking. It should provide role-based access for students and administrators, with features for application processing, status updates, and secure data management.""")
    doc.add_page_break()

    # Introduction
    create_title(doc, "2. Introduction", size=14)
    add_content(doc, """The College Admission Portal is a web-based Java application designed to modernize and streamline the college admission process. This application provides a user-friendly interface for students to submit applications online and for administrators to manage the admission process. The system is built using Java Servlets and JSP for the web interface and implements MVC architecture for efficient data management and processing.""")
    doc.add_paragraph()

    # Methodology
    create_title(doc, "3. Methodology", size=14)
    add_content(doc, "The project follows a structured development approach:")
    add_content(doc, "1. Analysis Phase:")
    add_content(doc, "   • Requirements gathering")
    add_content(doc, "   • System architecture design")
    add_content(doc, "   • Database design planning")
    add_content(doc, "2. Design Phase:")
    add_content(doc, "   • Class design (User, Application, Document)")
    add_content(doc, "   • Web interface design")
    add_content(doc, "   • Database schema design")
    add_content(doc, "3. Implementation Phase:")
    add_content(doc, "   • Frontend development using JSP and Bootstrap")
    add_content(doc, "   • Backend development using Java Servlets")
    add_content(doc, "   • Database implementation using MySQL")
    add_content(doc, "   • Testing and deployment")
    doc.add_page_break()

    # Implementation and Results
    create_title(doc, "4. Implementation and Results", size=14)
    add_content(doc, "The implementation includes:")
    add_content(doc, "1. Core Classes:")
    
    # Add code snippets
    add_content(doc, "a) User.java:")
    user_code = """public class User {
    private int id;
    private String email;
    private String password;
    private String role;
    private Timestamp createdAt;

    // Constructor
    public User(String email, String password, String role) {
        this.email = email;
        this.password = password;
        this.role = role;
    }

    // Getters and Setters
    public int getId() { return id; }
    public void setId(int id) { this.id = id; }
    // ... other getters and setters
}"""
    add_code(doc, user_code)

    add_content(doc, "b) Application.java:")
    application_code = """public class Application {
    private int id;
    private int userId;
    private String firstName;
    private String lastName;
    private Date dateOfBirth;
    private String phone;
    private String address;
    private String program;
    private String status;
    private Timestamp createdAt;
    private Timestamp updatedAt;

    // Constructor and methods
    // ... implementation details
}"""
    add_code(doc, application_code)

    add_content(doc, "c) ApplicationServlet.java:")
    servlet_code = """@WebServlet("/application/*")
@MultipartConfig
public class ApplicationServlet extends HttpServlet {
    private ApplicationDAO applicationDAO;

    @Override
    public void init() throws ServletException {
        applicationDAO = new ApplicationDAO();
    }

    protected void doPost(HttpServletRequest request, 
            HttpServletResponse response) throws ServletException, IOException {
        // Implementation details
    }
}"""
    add_code(doc, servlet_code)

    # Results section
    create_title(doc, "Results:", size=12)
    add_content(doc, "The implementation successfully provides:")
    add_content(doc, "• Secure user authentication and authorization")
    add_content(doc, "• Online application submission")
    add_content(doc, "• Document upload functionality")
    add_content(doc, "• Real-time status tracking")
    add_content(doc, "• Admin dashboard for application management")
    add_content(doc, "• Input validation and error handling")
    add_content(doc, "• Responsive web interface")
    add_content(doc, "• Efficient database operations")
    doc.add_page_break()

    # Outcome
    create_title(doc, "5. Outcome of the Project", size=14)
    add_content(doc, "The project successfully achieved:")
    add_content(doc, "1. Functional Outcomes:")
    add_content(doc, "   • Automated admission process")
    add_content(doc, "   • Online application submission")
    add_content(doc, "   • Document management")
    add_content(doc, "   • Status tracking system")
    add_content(doc, "2. Technical Outcomes:")
    add_content(doc, "   • Implementation of MVC architecture")
    add_content(doc, "   • Secure user authentication")
    add_content(doc, "   • Database integration")
    add_content(doc, "   • Responsive web design")
    add_content(doc, "3. Business Outcomes:")
    add_content(doc, "   • Streamlined admission process")
    add_content(doc, "   • Reduced paperwork")
    add_content(doc, "   • Improved efficiency")
    add_content(doc, "   • Better applicant experience")
    doc.add_page_break()

    # Conclusion
    create_title(doc, "6. Conclusion", size=14)
    add_content(doc, """The College Admission Portal successfully addresses the need for a modern, digital admission management system. The application demonstrates the practical application of Java web technologies and provides a user-friendly solution for both applicants and administrators. The implementation of MVC architecture, security features, and responsive design showcases the integration of technical skills with real-world requirements.""")
    doc.add_page_break()

    # References
    create_title(doc, "7. References", size=14)
    add_content(doc, "1. Java Servlet Programming by Jason Hunter")
    add_content(doc, "2. Head First Servlets and JSP by Kathy Sierra")
    add_content(doc, "3. MySQL Documentation")
    add_content(doc, "4. Bootstrap Documentation")
    add_content(doc, "5. Apache Tomcat Documentation")
    add_content(doc, "6. Maven Documentation")
    add_content(doc, "7. Web Application Security Guidelines")
    add_content(doc, "8. Database Design Principles")
    add_content(doc, "9. Web Development Best Practices")
    add_content(doc, "10. Software Engineering Standards")
    
    # Save the document
    doc.save('College_Admission_Portal_Report.docx')

if __name__ == "__main__":
    create_report()
    print("Report has been generated successfully!") 